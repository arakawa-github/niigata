import os
from flask import Flask, render_template, request,send_from_directory
from werkzeug.utils import secure_filename

from docx import Document
import pandas as pd
import re

app = Flask(__name__)

# アップロードフォルダ設定
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
DOWNLOAD_FOLDER = "download"
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)
app.config["DOWNLOAD_FOLDER"] = DOWNLOAD_FOLDER
# 許可する拡張子
ALLOWED_EXTENSIONS = {"docx"}

def wordFunc(doc_path):
    # Wordファイルを開く
    #doc_path = "ピッキングリスト.docx"
    doc = Document(doc_path)

    # データを格納するリスト
    data = []
    current_serial_number = None  # 連番を保持する変数
    tables = iter(doc.tables)  # すべての表を順番に処理するイテレータ

    # Word内のテキストを解析
    for para in doc.paragraphs:
        text = para.text.strip()

        # 「連番 H318」のようなフォーマットから H318 部分を抽出
        match = re.search(r"連番\s*([A-Za-z0-9]+)", text)
        if match:
            current_serial_number = match.group(1)  # 連番を更新
        
            # 次の表を取得（連番がある段落の後に表がある前提）
            try:
                table = next(tables)  # 次の表を取得
                table_data = []

                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                
                    # データが空白行でない場合のみ処理
                    if any(row_data):
                        table_data.append(row_data)

                # **表の最初の行（列名）を削除**
                if table_data:
                    table_data.pop(0)  

                # **処理後の表データを追加**
                for row in table_data:
                    if current_serial_number:  # 連番が取得できている場合のみ
                        data.append([current_serial_number] + row)

            except StopIteration:
               pass  # 表がもうない場合はスキップ




    # DataFrameを作成（最初の列は「連番」、残りは表のデータ）
    df = pd.DataFrame(data)

    # 4列目（インデックス3）、9列目（インデックス8）、10列目（インデックス9）、11列目（インデックス10）を削除
    df = df.drop(df.columns[[3, 7, 8, 9, 10]], axis=1)

    # 列名を設定
    df.columns = ["連番", "棚番", "品目コード", "指示数", "納入日", "品名"]
    #
    # 新しい列を追加（空の値で初期化）
    df["済み"] = ""
    df["担当者"] = ""
    df["ピッキング日"] = ""
    df["ID_XX"] = ""
    # Excelファイルに保存
    downloadName="ピッキングリスト.xlsx"
    df.to_excel("download/ピッキングリスト.xlsx", index=False)
    print("変換完了！")               
    
    return  downloadName               

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "ファイルが選択されていません", 400

    file = request.files["file"]
    if file.filename == "":
        return "ファイルが選択されていません", 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        ##
        file.save(file_path)
        ##
        downloadName = wordFunc(file_path)
        #print(downloadName)
        ##
        return render_template("index.html", filename=downloadName)  # ファイル名を渡す
    
    return "許可されていないファイル形式です", 400
    
@app.route("/download/<filename>")
def download_file(filename):
    return send_from_directory(app.config["DOWNLOAD_FOLDER"], filename, as_attachment=True)
    
if __name__ == "__main__":
    app.run(debug=True)
